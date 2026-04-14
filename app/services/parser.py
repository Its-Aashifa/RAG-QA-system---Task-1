"""
Document parsing service.
Supports PDF (via PyMuPDF), plain TXT, and DOCX (via python-docx).

Design principle: each format has its own private parser function.
Adding a new format = add one entry to SUPPORTED_EXTENSIONS and one _parse_X function.
No changes needed anywhere else in the system.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def parse_document(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    if ext == ".txt":
        return _parse_txt(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)


def _parse_txt(path: Path) -> str:
    """Read plain text file with UTF-8 encoding, fallback to latin-1."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decode failed for {path.name}, falling back to latin-1")
        return path.read_text(encoding="latin-1")


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

    doc = fitz.open(str(path))
    pages_text = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages_text.append(f"[Page {page_num + 1}]\n{text}")

    doc.close()

    if not pages_text:
        raise ValueError(
            f"No extractable text found in '{path.name}'. "
            "The PDF may be image-based and require OCR."
        )

    return "\n\n".join(pages_text)


def _parse_docx(path: Path) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Extracts paragraphs AND table content.
    Tables are included because business DOCX files often store key data in them.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(path))
    sections = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # Table content
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_rows.append(row_text)
        if table_rows:
            sections.append(f"[Table {table_idx + 1}]\n" + "\n".join(table_rows))

    if not sections:
        raise ValueError(f"No extractable text found in '{path.name}'.")

    return "\n\n".join(sections)
